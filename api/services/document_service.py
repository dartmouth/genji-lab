# services/document_service.py

from typing import List, Optional, Dict, Any
from datetime import datetime
from fastapi import HTTPException, status
from sqlalchemy.orm import Session, joinedload
from sqlalchemy import select, func, delete, Integer
from io import BytesIO
import docx
import re
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from models.models import (
    Document as DocumentModel,
    DocumentCollection,
    DocumentElement,
    Annotation as AnnotationModel
)
from schemas.documents import (
    DocumentCreate,
    DocumentUpdate,
    DocumentPartialUpdate
)
from services.base_service import BaseService


class DocumentService(BaseService[DocumentModel]):
    """Service for document CRUD operations."""
    
    def __init__(self):
        super().__init__(DocumentModel)
    
    # ==================== Word Document Processing Utilities ====================
    
    def _extract_links(self, paragraph_text: str) -> tuple[str, List[Dict[str, Any]]]:
        """
        Process paragraphs with links using string operations.
        Returns (clean_text, links).
        """
        number_pattern = re.compile(r"(\d+\.\d+)")
        links = []
        
        # Look for the characteristic beginning of the link
        link_marker = "[link to original paragraph"
        if link_marker in paragraph_text:
            # Find the position of the link
            start_idx = paragraph_text.rfind(link_marker)
            
            # Extract the link part
            link = paragraph_text[start_idx:]
            
            # Clean the text by removing the link
            clean_text = paragraph_text[:start_idx].strip()
            
            # Extract section and paragraph numbers
            try:
                # Extract the numbers part (e.g., "1.1")
                numbers_part = re.search(number_pattern, link).group(0)
                section_str, paragraph_str = numbers_part.split(".")
                
                link_info = {
                    "textCollectionId": "original",
                    "hierarchy": {
                        "section": int(section_str),
                        "chapter": int(section_str),
                        "paragraph": int(paragraph_str),
                    },
                }
                
                links.append(link_info)
            except Exception as e:
                print(f"Error parsing link: {e}")
            
            return clean_text, links
        
        return paragraph_text, links
    
    def _get_text_format(self, paragraph) -> Dict[str, Any]:
        """
        Extract text formatting information from a paragraph.
        """
        # Initialize formatting details
        formatting = []
        
        # Check if paragraph has any runs
        if not paragraph.runs:
            return {
                "is_bold": False,
                "is_italic": False,
                "is_underlined": False,
                "formatting": [],
            }
        
        # Initialize flags to check if all runs have the same formatting
        all_bold = True
        all_italic = True
        all_underlined = True
        
        # Check each run's formatting
        for run in paragraph.runs:
            # Update flags based on each run's formatting
            if not run.bold:
                all_bold = False
            if not run.italic:
                all_italic = False
            if not run.underline:
                all_underlined = False
            
            # Add specific formatting information for this run if it has italic or underline
            if run.italic or run.underline:
                start = paragraph.text.find(run.text)
                end = start + len(run.text)
                fmt = {"start": start, "end": end, "type": []}
                if run.italic:
                    fmt["type"].append("italic")
                if run.underline:
                    fmt["type"].append("underlined")
                formatting.append(fmt)
        
        # Return the complete formatting information
        return {
            "is_bold": all_bold,
            "is_italic": all_italic,
            "is_underlined": all_underlined,
            "formatting": formatting,
        }
    
    def _process_indent(self, value) -> float:
        """
        Convert indent value to normalized units.
        """
        if value is None:
            return 0
        return value / 914400
    
    def _get_paragraph_format(self, paragraph) -> Dict[str, Any]:
        """
        Extract formatting information from a paragraph.
        """
        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
            WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
            WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "justify",
        }
        
        left_indent = paragraph.paragraph_format.left_indent
        right_indent = paragraph.paragraph_format.right_indent
        first_line_indent = paragraph.paragraph_format.first_line_indent
        alignment = paragraph.paragraph_format.alignment
        
        text_styles = self._get_text_format(paragraph)
        
        return {
            "left_indent": self._process_indent(left_indent),
            "right_indent": self._process_indent(right_indent),
            "first_line_indent": self._process_indent(first_line_indent),
            "alignment": "left" if alignment is None else alignment_map[alignment],
            "text_styles": text_styles,
        }
    
    def _extract_paragraphs(
        self,
        doc,
        text_collection_id: int,
        document_number: int
    ) -> List[Dict[str, Any]]:
        """
        Extract paragraphs from a Word document and convert to structured format.
        """
        json_results = []
        element_counter = 1
        
        for idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():
                clean_text, links = self._extract_links(paragraph.text)
                paragraph_json = {
                    "document_collection_id": text_collection_id,
                    "hierarchy": {
                        "document": document_number,
                        "element_order": element_counter,
                    },
                    "content": {
                        "text": clean_text,
                        "formatting": self._get_paragraph_format(paragraph),
                    },
                }
                json_results.append(paragraph_json)
                element_counter += 1
        
        return json_results
    
    # ==================== Helper Methods ====================
    
    def _verify_collection_exists(self, db: Session, collection_id: int) -> DocumentCollection:
        """
        Verify a document collection exists by ID.
        
        Raises HTTPException 404 if not found.
        """
        collection = db.execute(
            select(DocumentCollection).filter(DocumentCollection.id == collection_id)
        ).scalar_one_or_none()
        
        if not collection:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Document collection with ID {collection_id} not found"
            )
        return collection
    
    def _get_document_by_id(
        self,
        db: Session,
        document_id: int,
        with_collection: bool = False
    ) -> DocumentModel:
        """
        Get a document by ID.
        
        Raises HTTPException 404 if not found.
        """
        query = select(DocumentModel).filter(DocumentModel.id == document_id)
        
        if with_collection:
            query = query.options(joinedload(DocumentModel.collection))
        
        document = db.execute(query).scalar_one_or_none()
        
        if document is None:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
        
        return document
    
    def _check_duplicate_title(
        self,
        db: Session,
        title: str,
        collection_id: int,
        exclude_id: Optional[int] = None
    ) -> None:
        """
        Check if a document with the given title already exists in the collection.
        
        Raises HTTPException 400 if duplicate found.
        """
        query = select(DocumentModel).filter(
            DocumentModel.document_collection_id == collection_id,
            func.lower(DocumentModel.title) == title.lower().strip()
        )
        
        if exclude_id is not None:
            query = query.filter(DocumentModel.id != exclude_id)
        
        existing = db.execute(query).scalar_one_or_none()
        
        if existing:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Document name already exists in this collection"
            )
    
    def _get_element_count(self, db: Session, document_id: int) -> int:
        """Get the count of elements for a document."""
        return db.execute(
            select(func.count())
            .select_from(DocumentElement)
            .filter(DocumentElement.document_id == document_id)
        ).scalar_one()
    
    def _get_scholarly_annotation_count(self, db: Session, document_id: int) -> int:
        """Get the count of scholarly annotations for a document."""
        return db.execute(
            select(func.count())
            .select_from(AnnotationModel)
            .filter(
                AnnotationModel.document_id == document_id,
                AnnotationModel.motivation.in_(['scholarly', 'highlighting', 'bookmarking', 'classifying'])
            )
        ).scalar_one()
    
    def _get_comment_count(self, db: Session, document_id: int) -> int:
        """Get the count of comments for a document."""
        return db.execute(
            select(func.count())
            .select_from(AnnotationModel)
            .filter(
                AnnotationModel.document_id == document_id,
                AnnotationModel.motivation == 'commenting'
            )
        ).scalar_one()
    
    def _cascade_delete_document_content(
        self,
        db: Session,
        document_id: int
    ) -> None:
        """Delete all elements and annotations for a document."""
        element_ids = db.execute(
            select(DocumentElement.id).filter(DocumentElement.document_id == document_id)
        ).scalars().all()
        
        if element_ids:
            db.execute(
                delete(AnnotationModel).where(
                    AnnotationModel.document_element_id.in_(element_ids)
                )
            )
            db.execute(
                delete(DocumentElement).where(DocumentElement.document_id == document_id)
            )
    
    def _cascade_delete_documents_content(
        self,
        db: Session,
        document_ids: List[int]
    ) -> None:
        """Delete all elements and annotations for multiple documents."""
        if not document_ids:
            return
        
        element_ids = db.execute(
            select(DocumentElement.id).filter(DocumentElement.document_id.in_(document_ids))
        ).scalars().all()
        
        if element_ids:
            db.execute(
                delete(AnnotationModel).where(
                    AnnotationModel.document_element_id.in_(element_ids)
                )
            )
            db.execute(
                delete(DocumentElement).where(DocumentElement.document_id.in_(document_ids))
            )
    
    # ==================== CRUD Operations ====================
    
    def create(
        self,
        db: Session,
        document: DocumentCreate
    ) -> DocumentModel:
        """
        Create a new document.
        
        Raises HTTPException 404 if collection not found.
        Raises HTTPException 400 if title already exists in collection.
        """
        self._verify_collection_exists(db, document.document_collection_id)
        self._check_duplicate_title(db, document.title, document.document_collection_id)
        
        db_document = DocumentModel(**document.model_dump())
        
        db.add(db_document)
        db.commit()
        db.refresh(db_document)
        
        return db_document
    
    def get_by_id(
        self,
        db: Session,
        document_id: int
    ) -> DocumentModel:
        """
        Get a specific document by ID with collection information and element count.
        
        Raises HTTPException 404 if not found.
        """
        document = self._get_document_by_id(db, document_id, with_collection=True)
        
        elements_count = self._get_element_count(db, document_id)
        setattr(document, "elements_count", elements_count)
        
        return document
    
    def list(
        self,
        db: Session,
        skip: int = 0,
        limit: int = 100,
        title: Optional[str] = None,
        collection_id: Optional[int] = None
    ) -> List[DocumentModel]:
        """Get documents with optional filtering and pagination."""
        query = select(DocumentModel)
        
        if title:
            query = query.filter(DocumentModel.title.ilike(f"%{title}%"))
        if collection_id:
            query = query.filter(DocumentModel.document_collection_id == collection_id)
        
        query = query.offset(skip).limit(limit)
        
        return db.execute(query).scalars().all()
    
    def update(
        self,
        db: Session,
        document_id: int,
        document: DocumentUpdate
    ) -> DocumentModel:
        """
        Full update of a document.
        
        Raises HTTPException 404 if document or collection not found.
        Raises HTTPException 400 if title already exists in collection.
        """
        db_document = self._get_document_by_id(db, document_id)
        
        if document.document_collection_id:
            self._verify_collection_exists(db, document.document_collection_id)
        
        if document.title:
            collection_id = document.document_collection_id or db_document.document_collection_id
            self._check_duplicate_title(db, document.title, collection_id, exclude_id=document_id)
        
        update_data = document.model_dump(exclude_unset=True)
        for key, value in update_data.items():
            setattr(db_document, key, value)
        
        db.commit()
        db.refresh(db_document)
        
        return db_document
    
    def partial_update(
        self,
        db: Session,
        document_id: int,
        document: DocumentPartialUpdate
    ) -> DocumentModel:
        """
        Partial update of a document.
        
        Raises HTTPException 404 if document or collection not found.
        Raises HTTPException 400 if title already exists in collection.
        """
        db_document = self._get_document_by_id(db, document_id)
        
        if document.document_collection_id:
            self._verify_collection_exists(db, document.document_collection_id)
        
        if document.title:
            collection_id = document.document_collection_id or db_document.document_collection_id
            self._check_duplicate_title(db, document.title, collection_id, exclude_id=document_id)
        
        update_data = document.model_dump(exclude_unset=True, exclude_none=True)
        for key, value in update_data.items():
            setattr(db_document, key, value)
        
        db.commit()
        db.refresh(db_document)
        
        return db_document
    
    def delete(
        self,
        db: Session,
        document_id: int,
        force: bool = True
    ) -> None:
        """
        Delete a document.
        
        If force=True, cascades delete through elements and annotations.
        
        Raises HTTPException 404 if not found.
        """
        db_document = self._get_document_by_id(db, document_id)
        
        if force:
            self._cascade_delete_document_content(db, document_id)
        
        db.delete(db_document)
        db.commit()
    
    def bulk_delete(
        self,
        db: Session,
        document_ids: List[int],
        force: bool = True
    ) -> None:
        """
        Bulk delete multiple documents.
        
        Raises HTTPException 400 if no IDs provided.
        Raises HTTPException 404 if any documents not found.
        """
        if not document_ids:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No document IDs provided"
            )
        
        existing_documents = db.execute(
            select(DocumentModel).filter(DocumentModel.id.in_(document_ids))
        ).scalars().all()
        
        existing_ids = [doc.id for doc in existing_documents]
        missing_ids = set(document_ids) - set(existing_ids)
        
        if missing_ids:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Documents not found: {list(missing_ids)}"
            )
        
        if force:
            self._cascade_delete_documents_content(db, document_ids)
        
        db.execute(delete(DocumentModel).where(DocumentModel.id.in_(document_ids)))
        db.commit()
    
    # ==================== Element Operations ====================
    
    def get_elements(
        self,
        db: Session,
        document_id: int,
        skip: int = 0,
        limit: int = 100
    ) -> List[Dict[str, Any]]:
        """
        Get all elements for a specific document.
        
        Raises HTTPException 404 if document not found.
        """
        self._get_document_by_id(db, document_id)
        
        elements = db.execute(
            select(DocumentElement)
            .filter(DocumentElement.document_id == document_id)
            .order_by(DocumentElement.hierarchy["element_order"].astext.cast(Integer))
            .offset(skip)
            .limit(limit)
        ).scalars().all()
        
        return [
            {
                "id": element.id,
                "document_id": element.document_id,
                "created": element.created,
                "modified": element.modified,
                "hierarchy": element.hierarchy,
                "content": element.content
            }
            for element in elements
        ]
    
    # ==================== Collection Operations ====================
    
    def get_by_collection_with_stats(
        self,
        db: Session,
        collection_id: int,
        skip: int = 0,
        limit: int = 100
    ) -> List[Dict[str, Any]]:
        """
        Get documents in a collection with annotation statistics.
        
        Raises HTTPException 404 if collection not found.
        """
        self._verify_collection_exists(db, collection_id)
        
        documents = db.execute(
            select(DocumentModel)
            .filter(DocumentModel.document_collection_id == collection_id)
            .offset(skip)
            .limit(limit)
        ).scalars().all()
        
        result = []
        for document in documents:
            scholarly_count = self._get_scholarly_annotation_count(db, document.id)
            comment_count = self._get_comment_count(db, document.id)
            element_count = self._get_element_count(db, document.id)
            
            result.append({
                "id": document.id,
                "title": document.title,
                "description": document.description,
                "created": document.created,
                "modified": document.modified,
                "scholarly_annotation_count": scholarly_count,
                "comment_count": comment_count,
                "element_count": element_count,
                "total_annotation_count": scholarly_count + comment_count
            })
        
        return result
    
    # ==================== Word Document Import ====================
    
    def import_word_document(
        self,
        db: Session,
        collection_id: int,
        title: str,
        description: str,
        file_content: bytes,
        filename: str
    ) -> Dict[str, Any]:
        """
        Create a new document and import Word document content.
        
        Raises HTTPException 400 if file is not a .docx.
        Raises HTTPException 404 if collection not found.
        Raises HTTPException 500 on processing errors.
        """
        if not filename.endswith(".docx"):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="File must be a .docx document"
            )
        
        self._verify_collection_exists(db, collection_id)
        
        try:
            document_data = DocumentCreate(
                title=title,
                description=description,
                document_collection_id=collection_id
            )
            
            db_document = DocumentModel(**document_data.dict())
            db.add(db_document)
            db.flush()
            
            doc = docx.Document(BytesIO(file_content))
            paragraph_count = len(doc.paragraphs)
            
            paragraphs = self._extract_paragraphs(doc, collection_id, db_document.id)
            
            created_elements = []
            for element_data in paragraphs:
                db_element = DocumentElement(
                    document_id=db_document.id,
                    content=element_data.get("content", {}),
                    hierarchy=element_data.get("hierarchy", 0),
                    created=datetime.now(),
                    modified=datetime.now()
                )
                db.add(db_element)
                created_elements.append(db_element)

            # Commit the entire transaction
            db.commit()

            # Refresh to get all IDs
            db.refresh(db_document)
            for element in created_elements:
                db.refresh(element)

            return {
                "document": {
                    "id": db_document.id,
                    "title": db_document.title,
                    "description": db_document.description,
                    "document_collection_id": db_document.document_collection_id,
                    "created": db_document.created.isoformat(),
                    "modified": db_document.modified.isoformat(),
                },
                "import_results": {
                    "filename": filename,
                    "paragraph_count": paragraph_count,
                    "elements_created": len(created_elements),
                    "message": "Document created and Word content imported successfully",
                },
            }

        except Exception as e:
            db.rollback()
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error creating document and importing Word content: {str(e)}",
            )

document_service = DocumentService()