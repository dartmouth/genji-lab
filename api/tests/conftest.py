# tests/conftest.py
import os
import sys

# Set test environment variables
os.environ.setdefault("DB_SCHEMA", "test_schema")
os.environ.setdefault("SQLALCHEMY_DATABASE_URL", "sqlite:///:memory:")

# Add project root to Python path
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, project_root)

# Configure pytest-asyncio
pytest_plugins = ('pytest_asyncio',)

import pytest
from datetime import datetime
from sqlalchemy import (
    create_engine,
    Column,
    Integer,
    String,
    Text,
    ForeignKey,
    DateTime,
    Boolean,
    JSON,
    Table,
)
from sqlalchemy.orm import declarative_base, relationship, sessionmaker, Session
from sqlalchemy.pool import StaticPool


# ==================== Test-Specific Base and Models ====================

TestBase = declarative_base()

# Association table for group members (no schema)
test_group_members = Table(
    "group_members",
    TestBase.metadata,
    Column("group_id", Integer, ForeignKey("groups.id")),
    Column("user_id", Integer, ForeignKey("users.id")),
)

# Association table for user roles (no schema)
test_user_roles = Table(
    "user_roles",
    TestBase.metadata,
    Column("user_id", Integer, ForeignKey("users.id")),
    Column("role_id", Integer, ForeignKey("roles.id")),
)


class TestUser(TestBase):
    """Test-specific User model without PostgreSQL-specific features."""
    __tablename__ = "users"

    id = Column(Integer, primary_key=True, index=True)
    first_name = Column(String(255))
    last_name = Column(String(255))
    email = Column(String(255), unique=True, index=True)
    username = Column(String(255), unique=True, index=True)
    is_active = Column(Boolean, default=True)
    viewed_tutorial = Column(Boolean, default=False)
    user_metadata = Column(JSON)

    # Relationships
    annotations = relationship(
        "TestAnnotation",
        foreign_keys="TestAnnotation.creator_id",
        back_populates="creator"
    )
    groups = relationship(
        "TestGroup",
        secondary=test_group_members,
        back_populates="members"
    )
    roles = relationship(
        "TestRole",
        secondary=test_user_roles,
        back_populates="users"
    )
    password_auth = relationship("UserPasswordModel", back_populates="user", uselist=False)


class TestGroup(TestBase):
    """Test-specific Group model without PostgreSQL-specific features."""
    __tablename__ = "groups"

    id = Column(Integer, primary_key=True, index=True)
    name = Column(String(100))
    description = Column(String(255))
    created_at = Column(DateTime, default=datetime.now)
    created_by_id = Column(Integer, ForeignKey("users.id"))

    # Relationships
    members = relationship(
        "TestUser",
        secondary=test_group_members,
        back_populates="groups"
    )
    created_by = relationship("TestUser", foreign_keys=[created_by_id])


class TestRole(TestBase):
    """Test-specific Role model without PostgreSQL-specific features."""
    __tablename__ = "roles"

    id = Column(Integer, primary_key=True, index=True)
    name = Column(String(50), unique=True)
    description = Column(String(255))

    # Relationships
    users = relationship(
        "TestUser",
        secondary=test_user_roles,
        back_populates="roles"
    )


class UserPasswordModel(TestBase):
    """Test-specific UserPassword model without PostgreSQL-specific features."""
    __tablename__ = "user_passwords"

    id = Column(Integer, primary_key=True, index=True)
    user_id = Column(Integer, ForeignKey("users.id"), unique=True)
    hashed_password = Column(String(255), nullable=False)
    created_at = Column(DateTime, default=datetime.now)
    updated_at = Column(DateTime, default=datetime.now, onupdate=datetime.now)

    # Relationships
    user = relationship("TestUser", back_populates="password_auth")


class TestDocumentCollection(TestBase):
    """Test-specific DocumentCollection model without PostgreSQL-specific features."""
    __tablename__ = "document_collections"

    id = Column(Integer, primary_key=True, index=True)
    title = Column(String(255))
    visibility = Column(String(50))
    text_direction = Column(String(50))
    created = Column(DateTime, default=datetime.now)
    modified = Column(DateTime, default=datetime.now, onupdate=lambda: datetime.now())
    created_by_id = Column(Integer, ForeignKey("users.id"))
    modified_by_id = Column(Integer, ForeignKey("users.id"))
    owner_id = Column(Integer, ForeignKey("users.id"))
    language = Column(String(50))
    hierarchy = Column(JSON)  # JSON instead of JSONB
    collection_metadata = Column(JSON)  # JSON instead of JSONB
    display_order = Column(Integer, nullable=False, default=0)

    # Relationships
    created_by = relationship(
        "TestUser",
        foreign_keys=[created_by_id]
    )
    modified_by = relationship(
        "TestUser",
        foreign_keys=[modified_by_id]
    )
    owner = relationship(
        "TestUser",
        foreign_keys=[owner_id]
    )


class TestDocument(TestBase):
    """Test-specific Document model without PostgreSQL-specific features."""
    __tablename__ = "documents"

    id = Column(Integer, primary_key=True, index=True)
    title = Column(String(255))
    description = Column(Text)
    document_collection_id = Column(Integer, ForeignKey("document_collections.id"))
    created = Column(DateTime, default=datetime.now)
    modified = Column(DateTime, default=datetime.now)
    
    # Relationships
    collection = relationship("TestDocumentCollection")
    elements = relationship("TestDocumentElement", back_populates="document")


class TestDocumentElement(TestBase):
    """Test-specific DocumentElement model without PostgreSQL-specific features."""
    __tablename__ = "document_elements"

    id = Column(Integer, primary_key=True, index=True)
    document_id = Column(Integer, ForeignKey("documents.id"))
    element_type = Column(String(100))
    content = Column(JSON)  # Changed from Text to JSON for proper serialization
    hierarchy = Column(JSON)
    created = Column(DateTime, default=datetime.now)
    modified = Column(DateTime, default=datetime.now)
    
    # Relationships
    document = relationship("TestDocument", back_populates="elements")
    annotations = relationship(
        "TestAnnotation",
        primaryjoin="TestDocumentElement.id == foreign(TestAnnotation.document_element_id)",
        viewonly=True
    )


class TestSiteSettings(TestBase):
    """Test-specific SiteSettings model without PostgreSQL-specific features."""
    __tablename__ = "site_settings"

    id = Column(Integer, primary_key=True, index=True)
    site_title = Column(String(255))
    site_logo_enabled = Column(Boolean, default=False)
    updated_by_id = Column(Integer)
    updated_at = Column(DateTime)
    site_logo_data = Column(Text)
    site_logo_mime_type = Column(String(100))
    site_favicon_data = Column(Text)
    site_favicon_mime_type = Column(String(100))
    collection_metadata_schema = Column(JSON)


class TestAnnotation(TestBase):
    """Test-specific Annotation model without PostgreSQL-specific features."""
    __tablename__ = "annotations"

    id = Column(Integer, primary_key=True, index=True)
    document_collection_id = Column(Integer)
    document_id = Column(Integer)
    document_element_id = Column(Integer)
    creator_id = Column(Integer, ForeignKey("users.id"))
    owner_id = Column(Integer, ForeignKey("users.id"))
    classroom_id = Column(Integer, ForeignKey("groups.id"), nullable=True)

    type = Column(String(100))
    created = Column(DateTime)
    modified = Column(DateTime)
    generator = Column(String(255))
    generated = Column(DateTime)
    motivation = Column(String(100))

    body = Column(JSON)  # JSON instead of JSONB
    target = Column(JSON)  # JSON instead of JSONB

    status = Column(String(50))
    annotation_type = Column(String(100))
    context = Column(String(255))

    # Relationships
    creator = relationship(
        "TestUser",
        foreign_keys=[creator_id],
        back_populates="annotations"
    )
    classroom = relationship("TestGroup", foreign_keys=[classroom_id])


class CASConfigurationModel(TestBase):
    """Test-specific CASConfiguration model without PostgreSQL-specific features."""
    __tablename__ = "cas_configuration"

    id = Column(Integer, primary_key=True, index=True)
    enabled = Column(Boolean, default=False)

    # Core CAS settings
    server_url = Column(String(255), nullable=True)
    validation_endpoint = Column(String(100), default="/serviceValidate")
    protocol_version = Column(String(10), default="2.0")

    # XML parsing configuration
    xml_namespace = Column(String(255), default="http://www.yale.edu/tp/cas")

    # Attribute mapping (JSON instead of JSONB)
    attribute_mapping = Column(
        JSON,
        default={
            "username": "netid",
            "email": "email",
            "first_name": "givenName",
            "last_name": "sn",
            "full_name": "name",
        },
    )

    # Username extraction patterns
    username_patterns = Column(
        JSON,
        default=[
            '<cas:attribute name="{attr}" value="([^"]+)"',
            "<cas:{attr}>([^<]+)</cas:{attr}>",
            "<cas:user>([^<]+)</cas:user>",
        ],
    )

    # Additional metadata fields to extract
    metadata_attributes = Column(JSON, default=["uid", "netid", "did", "affil"])

    # Email handling
    email_domain = Column(String(255), nullable=True)
    email_format = Column(String(50), default="from_cas")

    # Display settings
    display_name = Column(String(100), default="CAS Login")
    
    # Audit fields
    updated_at = Column(DateTime, nullable=True)
    updated_by_id = Column(Integer, nullable=True)


# ==================== Database Fixtures ====================

@pytest.fixture(scope="function")
def engine():
    """Create an in-memory SQLite engine for each test."""
    engine = create_engine(
        "sqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    return engine


@pytest.fixture(scope="function")
def tables(engine):
    """Create all tables before each test, drop after."""
    TestBase.metadata.create_all(bind=engine)
    yield
    TestBase.metadata.drop_all(bind=engine)


@pytest.fixture(scope="function")
def db_session(engine, tables) -> Session:
    """Create a new database session for each test."""
    SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=engine)
    session = SessionLocal()
    
    yield session
    
    session.rollback()
    session.close()


# ==================== Model Alias Fixtures ====================

@pytest.fixture
def User():
    """Provide TestUser as User for tests."""
    return TestUser


@pytest.fixture
def Group():
    """Provide TestGroup as Group for tests."""
    return TestGroup


@pytest.fixture
def AnnotationModel():
    """Provide TestAnnotation as AnnotationModel for tests."""
    return TestAnnotation


@pytest.fixture
def DocumentCollectionModel():
    """Provide TestDocumentCollection as DocumentCollectionModel for tests."""
    return TestDocumentCollection


@pytest.fixture
def Role():
    """Provide TestRole as Role for tests."""
    return TestRole


@pytest.fixture
def UserPassword():
    """Provide TestUserPassword as UserPassword for tests."""
    return TestUserPassword


# ==================== User Fixtures ====================

@pytest.fixture
def test_user(db_session) -> TestUser:
    """Create a regular test user in the database."""
    user = TestUser(
        id=1,
        first_name="Test",
        last_name="User",
        email="test@example.com",
        username="testuser",
        is_active=True,
    )
    db_session.add(user)
    db_session.commit()
    db_session.refresh(user)
    return user


@pytest.fixture
def admin_user(db_session) -> TestUser:
    """Create an admin test user in the database."""
    user = TestUser(
        id=2,
        first_name="Admin",
        last_name="User",
        email="admin@example.com",
        username="adminuser",
        is_active=True,
    )
    db_session.add(user)
    db_session.commit()
    db_session.refresh(user)
    return user


@pytest.fixture
def verified_scholar_user(db_session) -> TestUser:
    """Create a verified scholar user in the database."""
    user = TestUser(
        id=3,
        first_name="Scholar",
        last_name="User",
        email="scholar@example.com",
        username="scholaruser",
        is_active=True,
    )
    db_session.add(user)
    db_session.commit()
    db_session.refresh(user)
    return user


# ==================== Role Fixtures ====================

@pytest.fixture
def test_role_user(db_session) -> TestRole:
    """Create a 'user' role in the database."""
    role = TestRole(
        id=1,
        name="user",
        description="Regular user"
    )
    db_session.add(role)
    db_session.commit()
    db_session.refresh(role)
    return role


@pytest.fixture
def test_role_admin(db_session) -> TestRole:
    """Create an 'admin' role in the database."""
    role = TestRole(
        id=2,
        name="admin",
        description="System administrator"
    )
    db_session.add(role)
    db_session.commit()
    db_session.refresh(role)
    return role


@pytest.fixture
def test_role_verified_scholar(db_session) -> TestRole:
    """Create a 'verified_scholar' role in the database."""
    role = TestRole(
        id=3,
        name="verified_scholar",
        description="Verified scholar"
    )
    db_session.add(role)
    db_session.commit()
    db_session.refresh(role)
    return role


# ==================== Classroom/Group Fixtures ====================

@pytest.fixture
def test_classroom(db_session, test_user) -> TestGroup:
    """Create a test classroom with the test user as a member."""
    classroom = TestGroup(
        id=1,
        name="Test Classroom",
        description="A test classroom",
        created_by_id=test_user.id,
    )
    db_session.add(classroom)
    db_session.commit()
    db_session.refresh(classroom)
    
    classroom.members.append(test_user)
    db_session.commit()
    
    return classroom


# ==================== Annotation Data Fixtures ====================

@pytest.fixture
def valid_body_data():
    """Valid body data matching the Pydantic Body schema."""
    return {
        "id": 1,
        "type": "TextualBody",
        "value": "This is a test annotation",
        "format": "text/plain",
        "language": "en"
    }


@pytest.fixture
def valid_text_target_data():
    """Valid text target with selector matching TextTarget schema."""
    return {
        "id": 1,
        "type": "TextTarget",
        "source": "document/element/1",
        "selector": {
            "type": "TextQuoteSelector",
            "value": "selected text",
            "refined_by": {
                "type": "TextPositionSelector",
                "start": 0,
                "end": 13
            }
        }
    }


@pytest.fixture
def valid_object_target_data():
    """Valid object target without selector matching ObjectTarget schema."""
    return {
        "id": 2,
        "type": "ObjectTarget",
        "source": "image/1"
    }


@pytest.fixture
def valid_annotation_create_data(valid_body_data, valid_text_target_data):
    """Valid annotation creation data matching AnnotationCreate schema."""
    return {
        "context": "http://www.w3.org/ns/anno.jsonld",
        "document_collection_id": 1,
        "document_id": 1,
        "document_element_id": 1,
        "creator_id": 1,
        "type": "Annotation",
        "motivation": "commenting",
        "generator": "test-app",
        "body": valid_body_data,
        "target": [valid_text_target_data]
    }


@pytest.fixture
def cas_config():
    """Sample CAS configuration for testing."""
    config = CASConfigurationModel(
        id=1,
        enabled=True,
        server_url="https://login.example.edu/cas",
        validation_endpoint="/serviceValidate",
        protocol_version="2.0",
        xml_namespace="http://www.yale.edu/tp/cas",
        attribute_mapping={
            "username": "netid",
            "email": "email",
            "first_name": "givenName",
            "last_name": "sn",
            "full_name": "name",
        },
        username_patterns=[
            '<cas:attribute name="{attr}" value="([^"]+)"',
            "<cas:{attr}>([^<]+)</cas:{attr}>",
            "<cas:user>([^<]+)</cas:user>",
        ],
        metadata_attributes=["uid", "netid", "did", "affil"],
        email_domain="dartmouth.edu",
        email_format="from_cas",
        display_name="CAS Login"
    )
    return config


# ==================== Database Annotation Fixtures ====================

@pytest.fixture
def test_annotation(
    db_session, 
    test_user, 
    valid_body_data, 
    valid_text_target_data
) -> TestAnnotation:
    """Create a test annotation in the database."""
    annotation = TestAnnotation(
        id=1,
        context="http://www.w3.org/ns/anno.jsonld",
        document_collection_id=1,
        document_id=1,
        document_element_id=1,
        creator_id=test_user.id,
        classroom_id=None,
        type="Annotation",
        motivation="commenting",
        generator="test-app",
        generated=datetime.now(),
        body=valid_body_data,
        target=[valid_text_target_data],
        status="active",
        annotation_type="comment",
        created=datetime.now(),
        modified=datetime.now(),
    )
    db_session.add(annotation)
    db_session.commit()
    db_session.refresh(annotation)
    return annotation


@pytest.fixture
def test_annotation_with_classroom(
    db_session, 
    test_user, 
    test_classroom,
    valid_body_data, 
    valid_text_target_data
) -> TestAnnotation:
    """Create a test annotation associated with a classroom."""
    annotation = TestAnnotation(
        id=2,
        context="http://www.w3.org/ns/anno.jsonld",
        document_collection_id=1,
        document_id=1,
        document_element_id=1,
        creator_id=test_user.id,
        classroom_id=test_classroom.id,
        type="Annotation",
        motivation="commenting",
        generator="test-app",
        generated=datetime.now(),
        body=valid_body_data,
        target=[valid_text_target_data],
        status="active",
        annotation_type="comment",
        created=datetime.now(),
        modified=datetime.now(),
    )
    db_session.add(annotation)
    db_session.commit()
    db_session.refresh(annotation)
    return annotation


@pytest.fixture
def multiple_test_annotations(
    db_session, 
    test_user,
    valid_body_data,
    valid_text_target_data
) -> list[TestAnnotation]:
    """Create multiple test annotations for pagination and filtering tests."""
    annotations = []
    
    for i in range(5):
        body_data = valid_body_data.copy()
        body_data["id"] = i + 10
        body_data["value"] = f"Test annotation {i}"
        
        target_data = valid_text_target_data.copy()
        target_data["id"] = i + 10
        
        annotation = TestAnnotation(
            id=i + 10,
            context="http://www.w3.org/ns/anno.jsonld",
            document_collection_id=1,
            document_id=1,
            document_element_id=i + 1,
            creator_id=test_user.id,
            classroom_id=None,
            type="Annotation",
            motivation="commenting" if i % 2 == 0 else "highlighting",
            generator="test-app",
            generated=datetime.now(),
            body=body_data,
            target=[target_data],
            status="active",
            annotation_type="comment",
            created=datetime.now(),
            modified=datetime.now(),
        )
        db_session.add(annotation)
        annotations.append(annotation)
    
    db_session.commit()
    for ann in annotations:
        db_session.refresh(ann)
    
    return annotations


@pytest.fixture
def annotation_with_multiple_targets(
    db_session,
    test_user,
    valid_body_data,
    valid_text_target_data,
    valid_object_target_data
) -> TestAnnotation:
    """Create an annotation with multiple targets for target operation tests."""
    annotation = TestAnnotation(
        id=100,
        context="http://www.w3.org/ns/anno.jsonld",
        document_collection_id=1,
        document_id=1,
        document_element_id=1,
        creator_id=test_user.id,
        classroom_id=None,
        type="Annotation",
        motivation="linking",
        generator="test-app",
        generated=datetime.now(),
        body=valid_body_data,
        target=[valid_text_target_data, valid_object_target_data],
        status="active",
        annotation_type="link",
        created=datetime.now(),
        modified=datetime.now(),
    )
    db_session.add(annotation)
    db_session.commit()
    db_session.refresh(annotation)
    return annotation


@pytest.fixture
def annotation_with_nested_targets(
    db_session,
    test_user,
    valid_body_data
) -> TestAnnotation:
    """Create an annotation with nested target arrays."""
    nested_targets = [
        [
            {
                "id": 201,
                "type": "TextTarget",
                "source": "doc/1",
                "selector": None
            },
            {
                "id": 202,
                "type": "TextTarget",
                "source": "doc/2",
                "selector": None
            }
        ]
    ]
    
    annotation = TestAnnotation(
        id=200,
        context="http://www.w3.org/ns/anno.jsonld",
        document_collection_id=1,
        document_id=1,
        document_element_id=1,
        creator_id=test_user.id,
        classroom_id=None,
        type="Annotation",
        motivation="linking",
        generator="test-app",
        generated=datetime.now(),
        body=valid_body_data,
        target=nested_targets,
        status="active",
        annotation_type="link",
        created=datetime.now(),
        modified=datetime.now(),
    )
    db_session.add(annotation)
    db_session.commit()
    db_session.refresh(annotation)
    return annotation


# ==================== Service Fixtures ====================

# Sequence counters for SQLite (PostgreSQL uses database sequences)
_body_id_counter = 0
_target_id_counter = 0


def reset_sequence_counters():
    """Reset sequence counters for test isolation."""
    global _body_id_counter, _target_id_counter
    _body_id_counter = 0
    _target_id_counter = 0


@pytest.fixture
def annotation_service(db_session, monkeypatch):
    """
    Create AnnotationService instance configured for SQLite testing.
    
    Patches ID generation methods to use simple counters instead of PostgreSQL sequences.
    """
    from services.annotation_service import AnnotationService
    
    # Reset counters for test isolation
    reset_sequence_counters()
    
    service = AnnotationService()
    
    # Patch the model to use TestAnnotation (service now uses self.model everywhere)
    service.model = TestAnnotation
    
    # Patch ID generation methods to work with SQLite
    def mock_generate_body_id(db: Session) -> int:
        global _body_id_counter
        _body_id_counter += 1
        return _body_id_counter
    
    def mock_generate_target_id(db: Session) -> int:
        global _target_id_counter
        _target_id_counter += 1
        return _target_id_counter
    
    monkeypatch.setattr(service, "generate_body_id", mock_generate_body_id)
    monkeypatch.setattr(service, "generate_target_id", mock_generate_target_id)
    
    return service


# ==================== Document Collection Fixtures ====================

@pytest.fixture
def test_document_collection(db_session, test_user) -> TestDocumentCollection:
    """Create a test document collection in the database."""
    collection = TestDocumentCollection(
        id=1,
        title="Test Collection",
        visibility="public",
        text_direction="ltr",
        language="en",
        created_by_id=test_user.id,
        owner_id=test_user.id,
        hierarchy={"type": "sequence", "elements": []},
        collection_metadata={"description": "A test collection"},
        display_order=0
    )
    db_session.add(collection)
    db_session.commit()
    db_session.refresh(collection)
    return collection


@pytest.fixture
def multiple_test_document_collections(
    db_session,
    test_user
) -> list[TestDocumentCollection]:
    """Create multiple test document collections."""
    collections = [
        TestDocumentCollection(
            id=10,
            title="Collection 1",
            visibility="public",
            text_direction="ltr",
            language="en",
            created_by_id=test_user.id,
            owner_id=test_user.id,
            display_order=10
        ),
        TestDocumentCollection(
            id=11,
            title="Collection 2",
            visibility="private",
            text_direction="rtl",
            language="ar",
            created_by_id=test_user.id,
            owner_id=test_user.id,
            display_order=5
        ),
        TestDocumentCollection(
            id=12,
            title="Collection 3",
            visibility="public",
            text_direction="ltr",
            language="es",
            created_by_id=test_user.id,
            owner_id=test_user.id,
            display_order=15
        )
    ]
    
    for collection in collections:
        db_session.add(collection)
    
    db_session.commit()
    
    for collection in collections:
        db_session.refresh(collection)
    
    return collections


@pytest.fixture
def test_document(db_session, test_document_collection) -> TestDocument:
    """Create a test document in the database."""
    document = TestDocument(
        id=1,
        title="Test Document",
        description="This is a test document",
        document_collection_id=test_document_collection.id
    )
    db_session.add(document)
    db_session.commit()
    db_session.refresh(document)
    return document


@pytest.fixture
def test_document_with_elements(db_session, test_document_collection) -> dict:
    """Create a test document with elements in the database."""
    document = TestDocument(
        id=2,
        title="Document with Elements",
        description="This document has elements",
        document_collection_id=test_document_collection.id
    )
    db_session.add(document)
    db_session.commit()
    db_session.refresh(document)
    
    # Create elements
    elements = []
    for i in range(3):
        element = TestDocumentElement(
            id=i + 1,
            document_id=document.id,
            element_type="paragraph",
            content=f"Element {i+1} content"
        )
        db_session.add(element)
        elements.append(element)
    
    db_session.commit()
    
    for element in elements:
        db_session.refresh(element)
    
    return {"document": document, "elements": elements}


@pytest.fixture
def test_site_settings(db_session) -> TestSiteSettings:
    """Create test site settings with empty metadata schema."""
    settings = TestSiteSettings(
        id=1,
        site_title="Test Site",
        collection_metadata_schema=[]  # Empty schema means no validation
    )
    db_session.add(settings)
    db_session.commit()
    db_session.refresh(settings)
    return settings


@pytest.fixture
def document_collection_service(db_session, monkeypatch):
    """
    Create DocumentCollectionService instance configured for SQLite testing.
    """
    import services.document_collection_service as dc_service_module
    from services.document_collection_service import DocumentCollectionService
    
    # Patch models in the service module's namespace (before creating service)
    monkeypatch.setattr(dc_service_module, 'User', TestUser)
    monkeypatch.setattr(dc_service_module, 'Document', TestDocument)
    monkeypatch.setattr(dc_service_module, 'DocumentElement', TestDocumentElement)
    monkeypatch.setattr(dc_service_module, 'AnnotationModel', TestAnnotation)
    
    # Need to patch SiteSettings in the models.models module since it's imported there
    import models.models
    monkeypatch.setattr(models.models, 'SiteSettings', TestSiteSettings)
    
    service = DocumentCollectionService()
    
    # Patch the model to use TestDocumentCollection
    service.model = TestDocumentCollection
    
    return service


@pytest.fixture
def document_element_service(db_session, monkeypatch):
    """
    Create DocumentElementService instance configured for SQLite testing.
    """
    import services.document_element_service as de_service_module
    from services.document_element_service import DocumentElementService
    
    # Patch models in the service module's namespace
    monkeypatch.setattr(de_service_module, 'DocumentElementModel', TestDocumentElement)
    monkeypatch.setattr(de_service_module, 'Document', TestDocument)
    monkeypatch.setattr(de_service_module, 'AnnotationModel', TestAnnotation)
    
    service = DocumentElementService()
    
    # Patch the model to use TestDocumentElement
    service.model = TestDocumentElement
    
    return service


# ==================== .docx Test Fixtures ====================

@pytest.fixture
def create_simple_docx():
    """Create a simple .docx file with plain text paragraphs."""
    from docx import Document
    from io import BytesIO
    
    def _create():
        doc = Document()
        doc.add_paragraph("First paragraph with plain text.")
        doc.add_paragraph("Second paragraph with more content.")
        doc.add_paragraph("Third paragraph to test multiple elements.")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    return _create


@pytest.fixture
def create_empty_docx():
    """Create an empty .docx file with no paragraphs."""
    from docx import Document
    from io import BytesIO
    
    def _create():
        doc = Document()
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    return _create


@pytest.fixture
def create_formatted_docx():
    """Create a .docx file with formatted text (bold, italic, links)."""
    from docx import Document
    from docx.shared import RGBColor
    from io import BytesIO
    
    def _create():
        doc = Document()
        
        # Paragraph with bold and italic
        p1 = doc.add_paragraph()
        run1 = p1.add_run("This is ")
        run2 = p1.add_run("bold text")
        run2.bold = True
        run3 = p1.add_run(" and this is ")
        run4 = p1.add_run("italic text")
        run4.italic = True
        
        # Paragraph with hyperlink
        p2 = doc.add_paragraph()
        p2.add_run("Check out ")
        hyperlink_run = p2.add_run("this link")
        hyperlink_run.font.color.rgb = RGBColor(0, 0, 255)
        hyperlink_run.font.underline = True
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    return _create


@pytest.fixture
def create_indented_docx():
    """Create a .docx file with indented paragraphs to test hierarchy."""
    from docx import Document
    from docx.shared import Pt
    from io import BytesIO
    
    def _create():
        doc = Document()
        
        # Top level paragraph
        p1 = doc.add_paragraph("Top level item")
        
        # Indented paragraph (level 1)
        p2 = doc.add_paragraph("First level indent")
        p2.paragraph_format.left_indent = Pt(36)  # 0.5 inch
        
        # More indented (level 2)
        p3 = doc.add_paragraph("Second level indent")
        p3.paragraph_format.left_indent = Pt(72)  # 1 inch
        
        # Back to level 1
        p4 = doc.add_paragraph("Back to first level")
        p4.paragraph_format.left_indent = Pt(36)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
    
    return _create
